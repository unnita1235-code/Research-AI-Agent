"""Deep Research agent: CLI, FastAPI server, and in-memory job queue."""

from __future__ import annotations

import argparse
import asyncio
import io
import json
import logging
import os
import sys
import time
import uuid
from pathlib import Path
from typing import Any

import uvicorn
from dotenv import load_dotenv
from fastapi import BackgroundTasks, Body, FastAPI, HTTPException, Query, Request
from fastapi.responses import HTMLResponse, RedirectResponse, Response, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field, field_validator

import job_store
from graph import run_research
from research_extras import (
    ask_feature_enabled,
    ask_report_async,
    counterarguments_async,
    counterarguments_enabled,
    executive_one_liner_async,
    executive_summary_enabled,
    suggest_topic_strings_async,
    suggest_topics_enabled,
    tldr_async,
    tldr_enabled,
    translate_enabled,
    translate_report_async,
)
from research_options import ResearchOptions
from state import ResearchState

APP_VERSION = "0.3.0"

BASE_DIR = Path(__file__).resolve().parent
load_dotenv(BASE_DIR / ".env")

logger = logging.getLogger(__name__)

# In-memory job store (optional SQLite mirror via ``job_store``).
_jobs: dict[str, dict[str, Any]] = {}
_MAX_TOPIC_LEN = 4_000


def _persist_job(job_id: str) -> None:
    j = _jobs.get(job_id)
    if not j or not job_store.enabled():
        return
    row = {
        "id": job_id,
        "topic": j.get("topic", ""),
        "status": j.get("status", "pending"),
        "report": j.get("report"),
        "error": j.get("error"),
        "created": j.get("created", time.time()),
        "events": j.get("events", []),
        "options": j.get("options"),
        "latest": j.get("latest"),
        "extracted_facts": j.get("extracted_facts", []),
    }
    job_store.upsert_row(row)


def _resolve_job(job_id: str) -> dict[str, Any] | None:
    if job_id in _jobs:
        return _jobs[job_id]
    if job_store.enabled():
        r = job_store.fetch_job(job_id)
        if r:
            _jobs[job_id] = r
            return r
    return None



def _env_ok_for_research() -> list[str]:
    mode = (os.environ.get("LLM_PROVIDER") or "").strip().lower() or "anthropic"
    if mode in ("heuristic", "ollama"):
        required: tuple[str, ...] = ("TAVILY_API_KEY",)
    else:
        required = ("ANTHROPIC_API_KEY", "TAVILY_API_KEY")
    return [k for k in required if not os.environ.get(k)]


def _ensure_env_or_raise() -> None:
    missing = _env_ok_for_research()
    if missing:
        raise HTTPException(
            status_code=503,
            detail=f"Server missing required env: {', '.join(missing)}. Set LLM_PROVIDER=heuristic to skip Anthropic.",
        )


class ResearchIn(BaseModel):
    topic: str = Field(min_length=1, max_length=_MAX_TOPIC_LEN)
    options: ResearchOptions = Field(default_factory=ResearchOptions)

    @field_validator("topic")
    @classmethod
    def _strip_topic(cls, v: str) -> str:
        return (v or "").strip()


class ResearchAccepted(BaseModel):
    job_id: str
    status: str
    poll_url: str
    stream_url: str


def _state_payload(state: ResearchState, label: str, step: int) -> dict[str, Any]:
    fr = state.get("final_report") or ""
    return {
        "type": "step",
        "step": step,
        "label": label,
        "search_queries": len(state.get("search_queries") or []),
        "source_results": len(state.get("source_results") or []),
        "extracted_facts": len(state.get("extracted_facts") or []),
        "final_report": fr,
        "final_report_len": len(fr),
    }


async def _run_research_job(job_id: str, topic: str, ro: ResearchOptions) -> None:
    job = _jobs.get(job_id)
    if not job:
        return
    job["status"] = "running"
    _persist_job(job_id)
    try:
        opt_dict = ro.model_dump()

        async def on_update(state: ResearchState, step: int, label: str) -> None:
            j = _jobs.get(job_id)
            if not j:
                return
            j["events"].append(_state_payload(state, label, step))
            j["latest"] = _state_payload(state, label, step)

        out = await run_research(
            topic,
            stream_progress=False,
            on_state_update=on_update,
            options=opt_dict,
        )
        job = _jobs.get(job_id)
        if not job:
            return
        st: ResearchState = out or {}  # type: ignore[assignment]
        job["status"] = "completed"
        job["report"] = (st or {}).get("final_report", "") or ""
        job["extracted_facts"] = (st or {}).get("extracted_facts") or []
        job["events"].append(
            {
                "type": "done",
                "status": "completed",
                "final_report": job["report"],
                "final_report_len": len(job["report"] or ""),
            }
        )
        _persist_job(job_id)
    except Exception as e:  # noqa: BLE001
        logger.exception("Research job %s failed", job_id)
        j = _jobs.get(job_id)
        if j:
            j["status"] = "failed"
            j["error"] = str(e)
            j["events"].append({"type": "done", "status": "failed", "error": str(e)})
        _persist_job(job_id)


app = FastAPI(
    title="Deep Research Agent",
    version=APP_VERSION,
)
app.mount(
    "/static",
    StaticFiles(directory=str(BASE_DIR / "static")),
    name="static",
)

# Simple per-IP window for ``POST /research`` (set RATE_LIMIT_PER_MIN=0 to disable).
_rate_hits: dict[str, list[float]] = {}


def _check_research_rate(request: Request) -> None:
    cap = int((os.environ.get("RATE_LIMIT_PER_MIN") or "0").strip() or "0")
    if cap <= 0:
        return
    ip = (request.client.host if request.client else "local") or "local"
    now = time.time()
    w = _rate_hits.setdefault(ip, [])
    w[:] = [t for t in w if now - t < 60.0]
    if len(w) >= cap:
        raise HTTPException(429, detail="Rate limit: too many research jobs per minute from this address.")
    w.append(now)

@app.get("/", response_class=HTMLResponse)
def root() -> RedirectResponse:
    return RedirectResponse(url="/static/index.html", status_code=302)


@app.get("/config")
def public_config() -> dict[str, Any]:
    p = (os.environ.get("LLM_PROVIDER") or "").strip().lower() or "anthropic"
    return {
        "llm_provider": p,
        "tavily_configured": bool((os.environ.get("TAVILY_API_KEY") or "").strip()),
        "anthropic_configured": bool((os.environ.get("ANTHROPIC_API_KEY") or "").strip()),
        "suggest_topics_enabled": suggest_topics_enabled(),
        "executive_summary_enabled": executive_summary_enabled(),
        "ask_enabled": ask_feature_enabled(),
        "counterarguments_enabled": counterarguments_enabled(),
        "tldr_enabled": tldr_enabled(),
        "translate_enabled": translate_enabled(),
        "job_db_enabled": job_store.enabled(),
        "rate_limit_per_min": int((os.environ.get("RATE_LIMIT_PER_MIN") or "0").strip() or "0"),
        "app_version": APP_VERSION,
    }


class SuggestIn(BaseModel):
    seed: str | None = Field(default=None, max_length=200)


class SuggestOut(BaseModel):
    topics: list[str]


class ExecutiveSummaryOut(BaseModel):
    summary: str


@app.post("/research/suggest-topics", response_model=SuggestOut)
async def post_suggest_topics(
    body: SuggestIn | None = Body(default=None),
) -> SuggestOut:
    """Return suggested research topics; uses Anthropic when configured, else a static list."""
    b = body or SuggestIn()
    topics = await suggest_topic_strings_async(b.seed)
    return SuggestOut(topics=topics)


class AskIn(BaseModel):
    question: str = Field(min_length=1, max_length=2_000)


class TranslateIn(BaseModel):
    target_lang: str = Field("Spanish", min_length=2, max_length=80)


class AskOut(BaseModel):
    answer: str


class TextBlockOut(BaseModel):
    text: str


def _build_docx_bytes(text: str, title: str) -> bytes:
    from docx import Document  # type: ignore[import-untyped]

    d = Document()
    d.add_heading((title or "Report")[:200], 0)
    for block in (text or "").split("\n\n"):
        d.add_paragraph(block.strip() or " ")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _build_html_page(title: str, body_md: str) -> str:
    t = (title or "Report").replace("&", "&amp;").replace("<", "&lt;")
    esc = (body_md or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return (
        f"<!DOCTYPE html><html lang=\"en\"><head><meta charset=\"utf-8\"/>"
        f"<title>{t}</title></head><body>"
        f'<pre style="white-space:pre-wrap;font-family:system-ui,Segoe UI,sans-serif">{esc}</pre>'
        f"</body></html>"
    )


@app.post("/research/{job_id}/ask", response_model=AskOut)
async def post_ask_report(job_id: str, body: AskIn) -> AskOut:
    if not ask_feature_enabled():
        raise HTTPException(404, "Feature disabled. Set ENABLE_ASK=1 in the server environment to enable Q&A.")
    job = _resolve_job(job_id)
    if not job or job.get("status") != "completed":
        raise HTTPException(400, "A completed report is required for this job.")
    report = (job.get("report") or "").strip()
    if not report:
        raise HTTPException(400, "Report is empty.")
    facts = [str(x) for x in (job.get("extracted_facts") or []) if str(x).strip()]
    if not facts:
        facts = [report[:5000]]
    try:
        ans = await ask_report_async(body.question.strip(), report=report, fact_lines=facts)
    except ValueError as e:
        raise HTTPException(400, str(e)) from e
    return AskOut(answer=ans)


@app.post("/research/{job_id}/counterarguments", response_model=TextBlockOut)
async def post_counterargs(job_id: str) -> TextBlockOut:
    if not counterarguments_enabled():
        raise HTTPException(404, "Feature disabled. Set ENABLE_COUNTERARGUMENTS=1.")
    job = _resolve_job(job_id)
    if not job or job.get("status") != "completed":
        raise HTTPException(400, "A completed report is required for this job.")
    report = (job.get("report") or "").strip()
    if not report:
        raise HTTPException(400, "Report is empty.")
    try:
        text = await counterarguments_async(report)
    except ValueError as e:
        raise HTTPException(503, str(e)) from e
    return TextBlockOut(text=text or "")


@app.post("/research/{job_id}/tldr", response_model=TextBlockOut)
async def post_tldr(job_id: str) -> TextBlockOut:
    if not tldr_enabled():
        raise HTTPException(404, "Feature disabled. Set ENABLE_TLDR=1.")
    job = _resolve_job(job_id)
    if not job or job.get("status") != "completed":
        raise HTTPException(400, "A completed report is required for this job.")
    report = (job.get("report") or "").strip()
    if not report:
        raise HTTPException(400, "Report is empty.")
    try:
        text = await tldr_async(report)
    except ValueError as e:
        raise HTTPException(503, str(e)) from e
    return TextBlockOut(text=text or "")


@app.post("/research/{job_id}/translate", response_model=TextBlockOut)
async def post_translate_report(job_id: str, body: TranslateIn) -> TextBlockOut:
    if not translate_enabled():
        raise HTTPException(404, "Feature disabled. Set ENABLE_TRANSLATE=1.")
    job = _resolve_job(job_id)
    if not job or job.get("status") != "completed":
        raise HTTPException(400, "A completed report is required for this job.")
    report = (job.get("report") or "").strip()
    if not report:
        raise HTTPException(400, "Report is empty.")
    try:
        text = await translate_report_async(report, body.target_lang)
    except ValueError as e:
        raise HTTPException(503, str(e)) from e
    return TextBlockOut(text=text or "")


@app.get("/research/{job_id}/export")
def export_research(
    job_id: str,
    file_format: str = Query("md", alias="format", description="md, html, or docx"),
) -> Any:
    """Download report as ``md`` (default), ``docx``, or ``html``."""
    job = _resolve_job(job_id)
    if not job:
        raise HTTPException(404, "Job not found")
    report = (job.get("report") or "").strip()
    if not report:
        raise HTTPException(400, "No report to export.")
    topic = (job.get("topic") or "report")[:80]
    fmt = (file_format or "md").lower().strip()
    if fmt in ("md", "markdown", "text"):
        return Response(
            content=report,
            media_type="text/markdown; charset=utf-8",
            headers={
                "Content-Disposition": f'attachment; filename="{_safe_filename(topic)}-research.md"'
            },
        )
    if fmt == "html":
        html = _build_html_page(topic, report)
        return Response(
            content=html,
            media_type="text/html; charset=utf-8",
            headers={
                "Content-Disposition": f'attachment; filename="{_safe_filename(topic)}-research.html"'
            },
        )
    if fmt == "docx":
        raw = _build_docx_bytes(report, topic)
        return Response(
            content=raw,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{_safe_filename(topic)}-research.docx"'
            },
        )
    raise HTTPException(400, "format must be md, html, or docx")


def _safe_filename(s: str) -> str:
    return "".join(c if c.isalnum() or c in "._- " else "_" for c in s)[:100].strip() or "report"


@app.post("/research/{job_id}/executive-summary", response_model=ExecutiveSummaryOut)
async def post_executive_summary(job_id: str) -> ExecutiveSummaryOut:
    if not executive_summary_enabled():
        raise HTTPException(404, "Feature disabled. Set ENABLE_EXECUTIVE_SUMMARY=1 in the server environment.")
    job = _resolve_job(job_id)
    if not job or job.get("status") != "completed":
        raise HTTPException(400, "A completed report is required for this job.")
    report = (job.get("report") or "").strip()
    if not report:
        raise HTTPException(400, "Report is empty.")
    try:
        text = await executive_one_liner_async(report)
    except ValueError as e:
        raise HTTPException(503, str(e)) from e
    except Exception as e:  # noqa: BLE001
        logger.exception("Executive summary failed for job %s", job_id)
        raise HTTPException(500, "Summary generation failed") from e
    return ExecutiveSummaryOut(summary=text)


@app.get("/jobs")
def list_job_rows(limit: int = 20) -> dict[str, Any]:
    """Recent jobs (from SQLite when ``ENABLE_JOB_DB`` is on)."""
    return {"jobs": job_store.list_jobs(limit=limit)}


@app.post("/research", response_model=ResearchAccepted, status_code=202)
async def start_research(
    body: ResearchIn, background_tasks: BackgroundTasks, request: Request
) -> ResearchAccepted:
    _check_research_rate(request)
    _ensure_env_or_raise()
    job_id = str(uuid.uuid4())
    _jobs[job_id] = {
        "id": job_id,
        "topic": body.topic,
        "options": body.options.model_dump(),
        "status": "pending",
        "report": None,
        "error": None,
        "events": [],
        "created": time.time(),
        "latest": None,
        "extracted_facts": [],
    }
    _persist_job(job_id)
    base = str(request.base_url).rstrip("/")
    background_tasks.add_task(_run_research_job, job_id, body.topic, body.options)
    return ResearchAccepted(
        job_id=job_id,
        status="pending",
        poll_url=f"{base}/research/{job_id}",
        stream_url=f"{base}/research/{job_id}/stream",
    )


@app.get("/research/{job_id}")
def get_research_status(job_id: str) -> dict[str, Any]:
    job = _resolve_job(job_id)
    if not job:
        raise HTTPException(404, "Job not found")
    return {
        "id": job["id"],
        "topic": job["topic"],
        "status": job["status"],
        "report": job.get("report"),
        "error": job.get("error"),
        "events": job.get("events", []),
        "latest": job.get("latest"),
        "options": job.get("options"),
        "extracted_facts": job.get("extracted_facts", []),
    }


@app.get("/research/{job_id}/stream")
async def stream_research(job_id: str) -> StreamingResponse:
    if not _resolve_job(job_id):
        raise HTTPException(404, "Job not found")

    async def event_gen() -> Any:
        sent = 0
        while True:
            job = _jobs.get(job_id)
            if not job:
                yield f"data: {json.dumps({'type': 'error', 'message': 'Job gone'})}\n\n"
                return
            evs: list[dict[str, Any]] = job.get("events", [])
            for i in range(sent, len(evs)):
                yield f"data: {json.dumps(evs[i])}\n\n"
            sent = len(evs)
            st = job.get("status", "")
            if st in ("completed", "failed"):
                return
            await asyncio.sleep(0.2)

    return StreamingResponse(
        event_gen(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


def _cli() -> int:
    ap = argparse.ArgumentParser(
        description="Deep research agent: CLI, or `serve` for FastAPI (see also: uvicorn main:app).",
    )
    ap.add_argument(
        "args",
        nargs="*",
        help="With `serve`: host/port. Otherwise: research topic, or -q / --serve flags.",
    )
    ap.add_argument(
        "-q",
        "--quiet",
        action="store_true",
        help="CLI only: no step logs on stdout.",
    )
    ap.add_argument(
        "--host",
        default=os.environ.get("HOST", "127.0.0.1"),
    )
    ap.add_argument(
        "--port",
        type=int,
        default=int(os.environ.get("PORT", "8000")),
    )
    u = ap.parse_args()

    if u.args and u.args[0] == "serve":
        h = u.host
        p = u.port
        if len(u.args) > 1:
            h = u.args[1]
        if len(u.args) > 2:
            p = int(u.args[2])
        print(f"Starting server at http://{h}:{p}/", file=sys.stderr, flush=True)
        uvicorn.run("main:app", host=h, port=p, reload=False)
        return 0

    if _env_ok_for_research():
        print("Missing required env. Copy .env.example. Set LLM_PROVIDER=heuristic to skip Anthropic.", file=sys.stderr)
        return 1

    topic = " ".join(u.args).strip() if u.args else ""
    if not topic:
        topic = "Outline of how retrieval-augmented generation (RAG) works for beginners"
    if not u.quiet:
        logging.basicConfig(level=logging.INFO, format="%(levelname)s %(name)s: %(message)s")

    async def _go() -> ResearchState:
        return await run_research(topic, stream_progress=not u.quiet, on_state_update=None)

    try:
        out = asyncio.run(_go())
    except Exception as e:  # noqa: BLE001
        print(f"Run failed: {e}", file=sys.stderr)
        return 1
    print((out or {}).get("final_report", "").strip() or "(empty report)")
    return 0


if __name__ == "__main__":
    raise SystemExit(_cli())
